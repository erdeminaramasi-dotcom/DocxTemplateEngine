from docx import Document
import json


class TemplateAnalyzer:

    def __init__(self, docx_path):
        self.doc = Document(docx_path)

    def analyze(self):

        result = {
            "tables": [],
            "paragraphs": []
        }

        # -------- TABLES --------

        for table_index, table in enumerate(self.doc.tables):

            table_info = {
                "table_index": table_index,
                "rows": []
            }

            for row_index, row in enumerate(table.rows):

                row_info = {
                    "row_index": row_index,
                    "cells": []
                }

                for cell_index, cell in enumerate(row.cells):

                    row_info["cells"].append({
                        "cell_index": cell_index,
                        "text": cell.text
                    })

                table_info["rows"].append(row_info)

            result["tables"].append(table_info)

        # -------- PARAGRAPHS --------

        for i, p in enumerate(self.doc.paragraphs):

            result["paragraphs"].append({
                "paragraph_index": i,
                "text": p.text
            })

        return result


if __name__ == "__main__":

    analyzer = TemplateAnalyzer("../templates/sample.docx")

    data = analyzer.analyze()

    print(json.dumps(data, indent=4, ensure_ascii=False))
